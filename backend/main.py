import os
import base64
import json
from datetime import datetime, date
from typing import Optional, List
from io import BytesIO

from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import func

import httpx

from database import (
    get_db,
    init_db,
    Job,
    ApplicationHistory,
    BaseResume,
    GeneratedResume,
    ATSAnalysis,
    TechFitAnalysis,
)
from models import (
    JobCreate,
    JobUpdate,
    Job as JobModel,
    JobStatus,
    RemoteType,
    BaseResumeCreate,
    BaseResume as BaseResumeModel,
    ResumeType,
    GeneratedResume as GeneratedResumeModel,
    ATSAnalysis as ATSAnalysisModel,
    TechFitAnalysis as TechFitAnalysisModel,
    ApplicationHistoryCreate,
    ApplicationHistory as ApplicationHistoryModel,
    DashboardStats,
    ResumeGenerateRequest,
)

AI_SERVICE_URL = os.getenv("AI_SERVICE_URL", "http://ai-service:8001")

app = FastAPI(title="JobSync API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def startup():
    init_db()


@app.get("/")
def root():
    return {"message": "JobSync API", "version": "1.0.0"}


@app.get("/health")
def health():
    return {"status": "healthy"}


@app.get("/api/dashboard/stats", response_model=DashboardStats)
def get_dashboard_stats(db: Session = Depends(get_db)):
    total_jobs = db.query(func.count(Job.id)).scalar()
    active_applications = (
        db.query(func.count(Job.id))
        .filter(
            Job.status.in_(
                ["saved", "applied", "phone_screen", "interview", "executive_call"]
            )
        )
        .scalar()
    )
    interviews = (
        db.query(func.count(Job.id))
        .filter(Job.status.in_(["interview", "executive_call"]))
        .scalar()
    )
    offers = db.query(func.count(Job.id)).filter(Job.status == "offered").scalar()

    return DashboardStats(
        total_jobs=total_jobs or 0,
        active_applications=active_applications or 0,
        interviews=interviews or 0,
        offers=offers or 0,
    )


@app.get("/api/jobs", response_model=List[JobModel])
def list_jobs(status: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(Job)
    if status:
        query = query.filter(Job.status == status)
    jobs = query.order_by(Job.updated_at.desc()).all()
    return jobs


@app.get("/api/jobs/{job_id}", response_model=JobModel)
def get_job(job_id: int, db: Session = Depends(get_db)):
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@app.post("/api/jobs", response_model=JobModel)
def create_job(job: JobCreate, db: Session = Depends(get_db)):
    db_job = Job(**job.dict())
    db.add(db_job)
    db.commit()
    db.refresh(db_job)
    return db_job


@app.put("/api/jobs/{job_id}", response_model=JobModel)
def update_job(job_id: int, job: JobUpdate, db: Session = Depends(get_db)):
    db_job = db.query(Job).filter(Job.id == job_id).first()
    if not db_job:
        raise HTTPException(status_code=404, detail="Job not found")

    update_data = job.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_job, key, value)

    db.commit()
    db.refresh(db_job)
    return db_job


@app.delete("/api/jobs/{job_id}")
def delete_job(job_id: int, db: Session = Depends(get_db)):
    db_job = db.query(Job).filter(Job.id == job_id).first()
    if not db_job:
        raise HTTPException(status_code=404, detail="Job not found")
    db.delete(db_job)
    db.commit()
    return {"message": "Job deleted successfully"}


@app.post("/api/jobs/{job_id}/parse-description")
async def parse_job_description(
    job_id: int, description: dict, db: Session = Depends(get_db)
):
    db_job = db.query(Job).filter(Job.id == job_id).first()
    if not db_job:
        raise HTTPException(status_code=404, detail="Job not found")

    job_text = description.get("description", "")
    db_job.raw_description = job_text

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                f"{AI_SERVICE_URL}/parse-job",
                json={"description": job_text},
                timeout=60.0,
            )
            response.raise_for_status()
            parsed_data = response.json()

            db_job.company = parsed_data.get("company", db_job.company)
            db_job.position = parsed_data.get("position", db_job.position)
            db_job.location = parsed_data.get("location", db_job.location)
            db_job.salary = parsed_data.get("salary", db_job.salary)
            remote_str = parsed_data.get("remote_type", "").lower()
            if remote_str in ["remote", "hybrid", "on-site"]:
                db_job.remote_type = remote_str
            db_job.requirements = parsed_data.get("requirements")
            db_job.nice_to_have = parsed_data.get("nice_to_have")
            db_job.responsibilities = parsed_data.get("responsibilities")
            db_job.keywords = parsed_data.get("keywords")
            db_job.credentials = parsed_data.get("credentials")

            db.commit()
            db.refresh(db_job)
            return parsed_data
        except httpx.HTTPError as e:
            db.commit()
            raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")


@app.get("/api/jobs/{job_id}/history", response_model=List[ApplicationHistoryModel])
def get_job_history(job_id: int, db: Session = Depends(get_db)):
    history = (
        db.query(ApplicationHistory)
        .filter(ApplicationHistory.job_id == job_id)
        .order_by(ApplicationHistory.created_at.desc())
        .all()
    )
    return history


@app.post("/api/jobs/{job_id}/history", response_model=ApplicationHistoryModel)
def create_history_entry(
    job_id: int, entry: ApplicationHistoryCreate, db: Session = Depends(get_db)
):
    db_entry = ApplicationHistory(job_id=job_id, **entry.dict())
    db.add(db_entry)
    db.commit()
    db.refresh(db_entry)

    db_job = db.query(Job).filter(Job.id == job_id).first()
    if db_job:
        db_job.status = entry.status
        db.commit()

    return db_entry


@app.get("/api/resumes", response_model=List[BaseResumeModel])
def list_resumes(file_type: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(BaseResume)
    if file_type:
        query = query.filter(BaseResume.file_type == file_type)
    resumes = query.order_by(BaseResume.created_at.desc()).all()
    return resumes


@app.get("/api/resumes/{resume_id}")
def get_resume(resume_id: int, db: Session = Depends(get_db)):
    resume = db.query(BaseResume).filter(BaseResume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    return {
        "id": resume.id,
        "filename": resume.filename,
        "file_type": resume.file_type,
        "text_content": resume.text_content,
        "created_at": resume.created_at.isoformat() if resume.created_at else None,
    }


@app.post("/api/resumes", response_model=BaseResumeModel)
async def upload_resume(
    file: UploadFile = File(...),
    file_type: str = Form(...),
    db: Session = Depends(get_db),
):
    content = await file.read()
    text_content = None

    if file.filename.lower().endswith(".txt"):
        text_content = content.decode("utf-8")
    elif file.filename.lower().endswith(".pdf"):
        try:
            from PyPDF2 import PdfReader

            pdf_reader = PdfReader(BytesIO(content))
            text_content = "\n".join(
                [
                    page.extract_text()
                    for page in pdf_reader.pages
                    if page.extract_text()
                ]
            )
        except Exception:
            text_content = ""
    elif file.filename.lower().endswith(".docx"):
        try:
            from docx import Document

            doc = Document(BytesIO(content))
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception:
            text_content = ""

    db_resume = BaseResume(
        filename=file.filename,
        content=content,
        content_type=file.content_type,
        file_type=file_type,
        text_content=text_content,
    )
    db.add(db_resume)
    db.commit()
    db.refresh(db_resume)
    return db_resume


@app.delete("/api/resumes/{resume_id}")
def delete_resume(resume_id: int, db: Session = Depends(get_db)):
    resume = db.query(BaseResume).filter(BaseResume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    db.delete(resume)
    db.commit()
    return {"message": "Resume deleted successfully"}


@app.get(
    "/api/jobs/{job_id}/generated-resumes", response_model=List[GeneratedResumeModel]
)
def list_generated_resumes(job_id: int, db: Session = Depends(get_db)):
    resumes = (
        db.query(GeneratedResume)
        .filter(GeneratedResume.job_id == job_id)
        .order_by(GeneratedResume.created_at.desc())
        .all()
    )
    return resumes


@app.get("/api/generated-resumes/{resume_id}", response_model=GeneratedResumeModel)
def get_generated_resume(resume_id: int, db: Session = Depends(get_db)):
    resume = db.query(GeneratedResume).filter(GeneratedResume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Generated resume not found")
    return resume


@app.post("/api/generate-resume")
async def generate_resume(
    request: ResumeGenerateRequest, db: Session = Depends(get_db)
):
    job = db.query(Job).filter(Job.id == request.job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    example_resume = None
    if request.example_resume_id:
        example_resume = (
            db.query(BaseResume)
            .filter(BaseResume.id == request.example_resume_id)
            .first()
        )

    template_resume = None
    if request.template_resume_id:
        template_resume = (
            db.query(BaseResume)
            .filter(BaseResume.id == request.template_resume_id)
            .first()
        )

    async with httpx.AsyncClient() as client:
        try:
            payload = {
                "job_description": job.raw_description or "",
                "job_requirements": job.requirements,
                "job_keywords": job.keywords,
                "example_resume": example_resume.text_content
                if example_resume
                else None,
                "template_format": template_resume.text_content
                if template_resume
                else None,
            }

            response = await client.post(
                f"{AI_SERVICE_URL}/generate-resume", json=payload, timeout=120.0
            )
            response.raise_for_status()
            result = response.json()

            db_resume = GeneratedResume(
                job_id=request.job_id, content=result.get("resume", "")
            )
            db.add(db_resume)
            db.commit()
            db.refresh(db_resume)

            return {
                "id": db_resume.id,
                "content": db_resume.content,
                "created_at": db_resume.created_at.isoformat(),
            }
        except httpx.HTTPError as e:
            raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")


@app.put("/api/generated-resumes/{resume_id}", response_model=GeneratedResumeModel)
def update_generated_resume(
    resume_id: int, content: dict, db: Session = Depends(get_db)
):
    resume = db.query(GeneratedResume).filter(GeneratedResume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Generated resume not found")

    resume.content = content.get("content", resume.content)
    db.commit()
    db.refresh(resume)
    return resume


@app.get("/api/generated-resumes/{resume_id}/export")
def export_resume(resume_id: int, db: Session = Depends(get_db)):
    resume = db.query(GeneratedResume).filter(GeneratedResume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Generated resume not found")

    job = db.query(Job).filter(Job.id == resume.job_id).first()

    from docx import Document

    doc = Document()

    title = f"{job.company} - {job.position}" if job else "Generated Resume"
    doc.add_heading(title, 0)

    for paragraph in resume.content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = (
        f"resume_{job.company}_{job.position}.docx"
        if job
        else f"resume_{resume_id}.docx"
    )
    filename = filename.replace(" ", "_").lower()

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/api/jobs/{job_id}/analyze-ats", response_model=ATSAnalysisModel)
async def analyze_ats(job_id: int, request: dict, db: Session = Depends(get_db)):
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    resume_id = request.get("resume_id")
    resume_content = request.get("resume_content", "")

    if resume_id:
        gen_resume = (
            db.query(GeneratedResume).filter(GeneratedResume.id == resume_id).first()
        )
        if gen_resume:
            resume_content = gen_resume.content

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                f"{AI_SERVICE_URL}/analyze-ats",
                json={
                    "resume": resume_content,
                    "job_description": job.raw_description or "",
                    "keywords": job.keywords or "",
                },
                timeout=60.0,
            )
            response.raise_for_status()
            result = response.json()

            analysis = ATSAnalysis(
                job_id=job_id,
                resume_id=resume_id,
                parse_score=result.get("parse_score", 0.0),
                keyword_match=result.get("keyword_match", 0.0),
                search_relevance=result.get("search_relevance", 0.0),
                overall_score=result.get("overall_score", 0.0),
                issues=result.get("issues"),
                recommendations=result.get("recommendations"),
                keywords_found=result.get("keywords_found"),
                keywords_missing=result.get("keywords_missing"),
            )
            db.add(analysis)
            db.commit()
            db.refresh(analysis)

            return ATSAnalysisModel(
                id=analysis.id,
                job_id=analysis.job_id,
                resume_id=analysis.resume_id,
                parse_score=float(analysis.parse_score),
                keyword_match=float(analysis.keyword_match),
                search_relevance=float(analysis.search_relevance),
                overall_score=float(analysis.overall_score),
                issues=analysis.issues,
                recommendations=analysis.recommendations,
                keywords_found=analysis.keywords_found,
                keywords_missing=analysis.keywords_missing,
            )
        except httpx.HTTPError as e:
            raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")


@app.post("/api/jobs/{job_id}/analyze-tech-fit", response_model=TechFitAnalysisModel)
async def analyze_tech_fit(job_id: int, request: dict, db: Session = Depends(get_db)):
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    resume_id = request.get("resume_id")
    resume_content = request.get("resume_content", "")

    if resume_id:
        gen_resume = (
            db.query(GeneratedResume).filter(GeneratedResume.id == resume_id).first()
        )
        if gen_resume:
            resume_content = gen_resume.content

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                f"{AI_SERVICE_URL}/analyze-tech-fit",
                json={
                    "resume": resume_content,
                    "job_description": job.raw_description or "",
                    "requirements": job.requirements or "",
                },
                timeout=60.0,
            )
            response.raise_for_status()
            result = response.json()

            analysis = TechFitAnalysis(
                job_id=job_id,
                resume_id=resume_id,
                skill_match=result.get("skill_match", 0.0),
                experience_relevance=result.get("experience_relevance", 0.0),
                leadership_fit=result.get("leadership_fit", 0.0),
                strengths=result.get("strengths"),
                gaps=result.get("gaps"),
                recommendations=result.get("recommendations"),
            )
            db.add(analysis)
            db.commit()
            db.refresh(analysis)

            return TechFitAnalysisModel(
                id=analysis.id,
                job_id=analysis.job_id,
                resume_id=analysis.resume_id,
                skill_match=float(analysis.skill_match),
                experience_relevance=float(analysis.experience_relevance),
                leadership_fit=float(analysis.leadership_fit),
                strengths=analysis.strengths,
                gaps=analysis.gaps,
                recommendations=analysis.recommendations,
            )
        except httpx.HTTPError as e:
            raise HTTPException(status_code=500, detail=f"AI service error: {str(e)}")


@app.get("/api/jobs/{job_id}/ats-analyses", response_model=List[ATSAnalysisModel])
def get_ats_analyses(job_id: int, db: Session = Depends(get_db)):
    analyses = (
        db.query(ATSAnalysis)
        .filter(ATSAnalysis.job_id == job_id)
        .order_by(ATSAnalysis.created_at.desc())
        .all()
    )
    return [
        ATSAnalysisModel(
            id=a.id,
            job_id=a.job_id,
            resume_id=a.resume_id,
            parse_score=float(a.parse_score),
            keyword_match=float(a.keyword_match),
            search_relevance=float(a.search_relevance),
            overall_score=float(a.overall_score),
            issues=a.issues,
            recommendations=a.recommendations,
            keywords_found=a.keywords_found,
            keywords_missing=a.keywords_missing,
        )
        for a in analyses
    ]


@app.get(
    "/api/jobs/{job_id}/tech-fit-analyses", response_model=List[TechFitAnalysisModel]
)
def get_tech_fit_analyses(job_id: int, db: Session = Depends(get_db)):
    analyses = (
        db.query(TechFitAnalysis)
        .filter(TechFitAnalysis.job_id == job_id)
        .order_by(TechFitAnalysis.created_at.desc())
        .all()
    )
    return [
        TechFitAnalysisModel(
            id=a.id,
            job_id=a.job_id,
            resume_id=a.resume_id,
            skill_match=float(a.skill_match),
            experience_relevance=float(a.experience_relevance),
            leadership_fit=float(a.leadership_fit),
            strengths=a.strengths,
            gaps=a.gaps,
            recommendations=a.recommendations,
        )
        for a in analyses
    ]


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
